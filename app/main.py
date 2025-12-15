from datetime import datetime
import io
import json
from typing import Optional
from base64 import b64decode
from sqlalchemy import or_
import base64
from flask import Blueprint, render_template, request, redirect, url_for, abort, send_file
from flask_login import login_required, current_user
from docx import Document as DocxDocument

from . import db
from .decorators import roles_required
from .models import User, Document, AuditEvent, DocumentType, Role, SecureMessage

from .security import (
    ensure_user_keys,
    build_snapshot,
    streebog256,
    gost_sign,
    gost_verify,
    encrypt_for_recipient,
    decrypt_for_recipient,
    kuz_cmac,
    b64e,
    b64d,
)

main_bp = Blueprint("main", __name__)

# -------------------- Справочники --------------------

DOC_TYPE_LABELS = {
    DocumentType.LETTER: "Письмо",
    DocumentType.INSTRUCTION: "Инструкция",
    DocumentType.PACKET: "Пакет",
}

ROLE_TITLES = {
    Role.CLERK: "Делопроизводитель",
    Role.APPROVER: "Генеральный директор",
    Role.EXECUTOR: "Исполнитель",
    Role.PDZK: "ПДЗК",
    Role.AUDITOR: "Наблюдатель",
}

_role_admin = getattr(Role, "ADMIN", None)
if _role_admin is not None:
    ROLE_TITLES[_role_admin] = "Администратор"

ROLE_TO = {
    Role.CLERK: "Делопроизводителю",
    Role.APPROVER: "Генеральному директору",
    Role.EXECUTOR: "Исполнителю",
    Role.PDZK: "ПДЗК",
    Role.AUDITOR: "Наблюдателю",
}
if _role_admin is not None:
    ROLE_TO[_role_admin] = "Администратору"


# -------------------- Строгие шаги по типам документов --------------------

STAGE_FLOW = {
    DocumentType.LETTER: [
        ("let_received", "Принято и передано данному лицу (Ф6 1–10, Ф7 1–7)"),
        ("let_execution", "Исполнение (Ф1 1–9, Ф1 1–6, Ф7 12)"),
        ("let_filed", "В дело"),
        ("let_pdzk", "ПДЗК с присутствием лица, протокол"),
    ],

    DocumentType.INSTRUCTION: [
        ("ins_reg_1_1_7", "Регистрация (Ф1 1–7)"),
        ("ins_agreement", "Согласование"),
        ("ins_approval", "Утверждение"),
        ("ins_reg_1_8", "Фиксация утверждения (Ф1 8)"),
        ("ins_reg_5_1_7", "Выделенное хранение (Ф5 1–7)"),
        ("ins_reg_1_1_3", "Учетный шаг (Ф1 1–3)"),
        ("ins_familiarization", "Ознакомление (Ф1 4)"),
        ("ins_issue_new", "Издание новой инструкции"),
        ("ins_pdzk", "ПДЗК, протокол"),
    ],

    DocumentType.PACKET: [
        ("pack_reg_6_1_7", "Первая регистрация пакета (Ф6 1–7)"),
        ("pack_reg_7_1_10", "Регистрация документов пакета (Ф7 1–10)"),
        ("pack_reg_1_1_3", "Учет носителей (Ф1 1–3)"),
        ("pack_familiarization", "Ознакомление (Ф1 4)"),
        ("pack_storage_3y", "Документ находится на хранении 3 года"),
        ("pack_pdzk", "ПДЗК, протокол"),
        ("pack_open_act", "Открытое приложение: акт передачи (Ф5 8–10)"),
        ("pack_conf_decision", "Конфиденциальное приложение: решение"),
    ],
}

TERMINAL_STAGES = {
    "archived": "Передан в архив",
    "destroyed": "Уничтожен",
    "declassified": "Рассекречен",
    "deleted": "Удалён наблюдателем",
}

STAGE_LABELS = {k: v for flow in STAGE_FLOW.values() for k, v in flow}
STAGE_LABELS.update(TERMINAL_STAGES)

STAGE_LABELS.update({
    "created": "Создан (старый документ)",
    "agreement": "Согласование (старый документ)",
    "approval": "Утверждение (старый документ)",
    "familiarization": "Ознакомление (старый документ)",
    "execution": "Исполнение (старый документ)",
    "pdzk": "ПДЗК (старый документ)",
})


# -------------------- Действия --------------------

ACTIONS = {
    # -------- ПИСЬМО --------
    "let_start_execution": {
        "label": "Перевести в исполнение",
        "types": [DocumentType.LETTER],
        "from": ["let_received"],
        "to": "let_execution",
        "roles": [Role.CLERK],
        "to_role": Role.EXECUTOR,
    },
    "let_to_filed": {
        "label": "В дело",
        "types": [DocumentType.LETTER],
        "from": ["let_execution"],
        "to": "let_filed",
        "roles": [Role.EXECUTOR],
        "to_role": Role.CLERK,
    },
    "let_to_pdzk": {
        "label": "Назначить ПДЗК",
        "types": [DocumentType.LETTER],
        "from": ["let_filed"],
        "to": "let_pdzk",
        "roles": [Role.CLERK],
        "to_role": Role.PDZK,
    },

    # -------- ИНСТРУКЦИЯ --------
    "ins_to_agreement": {
        "label": "Отправить на согласование",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_reg_1_1_7"],
        "to": "ins_agreement",
        "roles": [Role.CLERK],
        "to_role": Role.APPROVER,
    },
    "ins_agree": {
        "label": "Согласовано",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_agreement"],
        "to": "ins_approval",
        "roles": [Role.APPROVER],
        "to_role": Role.APPROVER,
    },
    "ins_reject": {
        "label": "Вернуть на доработку",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_agreement", "ins_approval"],
        "to": "ins_reg_1_1_7",
        "roles": [Role.APPROVER],
        "to_role": Role.CLERK,
    },
    "ins_approve": {
        "label": "Утверждено (Ф1 8)",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_approval"],
        "to": "ins_reg_1_8",
        "roles": [Role.APPROVER],
        "to_role": Role.CLERK,
    },
    "ins_to_storage": {
        "label": "Передать в выделенное хранение (Ф5 1–7)",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_reg_1_8"],
        "to": "ins_reg_5_1_7",
        "roles": [Role.CLERK],
        "to_role": Role.CLERK,
    },
    "ins_to_reg_1_1_3": {
        "label": "Учетный шаг (Ф1 1–3)",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_reg_5_1_7"],
        "to": "ins_reg_1_1_3",
        "roles": [Role.CLERK],
        "to_role": Role.CLERK,
    },
    "ins_to_familiarization": {
        "label": "Отправить на ознакомление (Ф1 4)",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_reg_1_1_3"],
        "to": "ins_familiarization",
        "roles": [Role.CLERK],
        "to_role": Role.EXECUTOR,
    },
    "ins_familiarized": {
        "label": "Ознакомление завершено",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_familiarization"],
        "to": "ins_issue_new",
        "roles": [Role.EXECUTOR],
        "to_role": Role.CLERK,
    },
    "ins_issued_new": {
        "label": "Издана новая инструкция",
        "types": [DocumentType.INSTRUCTION],
        "from": ["ins_issue_new"],
        "to": "ins_pdzk",
        "roles": [Role.CLERK],
        "to_role": Role.PDZK,
    },

    # -------- ПАКЕТ --------
    "pack_register_docs": {
        "label": "Зарегистрировать документы пакета (Ф7 1–10)",
        "types": [DocumentType.PACKET],
        "from": ["pack_reg_6_1_7"],
        "to": "pack_reg_7_1_10",
        "roles": [Role.CLERK],
        "to_role": Role.CLERK,
    },
    "pack_register_media": {
        "label": "Учет носителей (Ф1 1–3)",
        "types": [DocumentType.PACKET],
        "from": ["pack_reg_7_1_10"],
        "to": "pack_reg_1_1_3",
        "roles": [Role.CLERK],
        "to_role": Role.CLERK,
    },
    "pack_to_familiarization": {
        "label": "Ознакомление (Ф1 4)",
        "types": [DocumentType.PACKET],
        "from": ["pack_reg_1_1_3"],
        "to": "pack_familiarization",
        "roles": [Role.CLERK],
        "to_role": Role.EXECUTOR,
    },
    "pack_to_storage": {
        "label": "Передать на хранение 3 года",
        "types": [DocumentType.PACKET],
        "from": ["pack_familiarization"],
        "to": "pack_storage_3y",
        "roles": [Role.EXECUTOR],
        "to_role": Role.CLERK,
    },
    "pack_fast_forward_3y": {
        "label": "Прошло 3 года (промотать)",
        "types": [DocumentType.PACKET],
        "from": ["pack_storage_3y"],
        "to": "pack_pdzk",
        "roles": [Role.CLERK],
        "to_role": Role.PDZK,
    },
    "pack_pdzk_to_open_act": {
        "label": "Открытое приложение: акт передачи (Ф5 8–10)",
        "types": [DocumentType.PACKET],
        "from": ["pack_pdzk"],
        "to": "pack_open_act",
        "roles": [Role.PDZK],
        "to_role": Role.CLERK,
    },
    "pack_pdzk_to_conf_decision": {
        "label": "Конфиденциальное приложение: решение",
        "types": [DocumentType.PACKET],
        "from": ["pack_pdzk"],
        "to": "pack_conf_decision",
        "roles": [Role.PDZK],
        "to_role": Role.CLERK,
    },

    # -------- ФИНАЛЬНЫЕ РЕШЕНИЯ --------
    "archive": {
        "label": "Передать в архив",
        "types": [DocumentType.LETTER, DocumentType.INSTRUCTION, DocumentType.PACKET],
        "from": ["let_pdzk", "ins_pdzk", "pack_open_act", "pack_conf_decision"],
        "to": "archived",
        "roles": [Role.PDZK],
    },
    "destroy": {
        "label": "Уничтожить",
        "types": [DocumentType.LETTER, DocumentType.INSTRUCTION, DocumentType.PACKET],
        "from": ["let_pdzk", "ins_pdzk", "pack_open_act", "pack_conf_decision"],
        "to": "destroyed",
        "roles": [Role.PDZK],
    },
    "declassify": {
        "label": "Рассекретить",
        "types": [DocumentType.LETTER, DocumentType.INSTRUCTION, DocumentType.PACKET],
        "from": ["let_pdzk", "ins_pdzk", "pack_open_act", "pack_conf_decision"],
        "to": "declassified",
        "roles": [Role.PDZK],
    },
}


# -------------------- JSON-хелперы --------------------

def load_payload(d: Document) -> dict:
    try:
        return json.loads(d.content_json or "{}")
    except Exception:
        return {}


def save_payload(d: Document, payload: dict) -> None:
    d.content_json = json.dumps(payload, ensure_ascii=False)


def set_assigned_role(d: Document, role: Optional[str], payload: Optional[dict] = None) -> None:
    payload = payload if payload is not None else load_payload(d)
    if role:
        payload["_assigned_to_role"] = role
    else:
        payload.pop("_assigned_to_role", None)
    save_payload(d, payload)
    setattr(d, "assigned_to_role", role)


def attach_runtime_assigned(d: Document, payload: Optional[dict] = None) -> None:
    payload = payload if payload is not None else load_payload(d)
    setattr(d, "assigned_to_role", payload.get("_assigned_to_role"))


# -------------------- Вспомогательные функции --------------------

def initial_stage(doc_type: str) -> str:
    if doc_type == DocumentType.LETTER:
        return "let_received"
    if doc_type == DocumentType.INSTRUCTION:
        return "ins_reg_1_1_7"
    if doc_type == DocumentType.PACKET:
        return "pack_reg_6_1_7"
    return "created"


def is_terminal(stage: str) -> bool:
    return stage in TERMINAL_STAGES


def calc_progress(doc_type: str, stage: str):
    flow = STAGE_FLOW.get(doc_type, [])
    total = len(flow)

    if is_terminal(stage):
        return 100, total, total, TERMINAL_STAGES[stage]

    keys = [k for k, _ in flow]
    if stage not in keys or total == 0:
        return 0, 0, 0, STAGE_LABELS.get(stage, stage)

    idx = keys.index(stage)
    step_now = idx + 1
    percent = round(step_now / total * 100)
    return percent, step_now, total, STAGE_LABELS.get(stage, stage)


def available_actions(doc_type: str, stage: str, role: str, assigned_to_role: Optional[str]):
    if assigned_to_role and assigned_to_role != role:
        return []

    res = []
    for action_id, meta in ACTIONS.items():
        if doc_type not in meta.get("types", []):
            continue
        if stage not in meta.get("from", []):
            continue
        if role not in meta.get("roles", []):
            continue

        to_role = meta.get("to_role")
        to_who = ROLE_TO.get(to_role) if to_role else None

        res.append({
            "id": action_id,
            "label": meta.get("label", action_id),
            "to": meta.get("to"),
            "to_who": to_who,
        })
    return res


def make_reg_number(doc_type: str) -> str:
    prefix = {
        DocumentType.LETTER: "LT",
        DocumentType.INSTRUCTION: "IN",
        DocumentType.PACKET: "PK",
    }.get(doc_type, "DOC")

    year = datetime.now().year
    last = Document.query.filter_by(doc_type=doc_type).count() + 1
    return f"{prefix}-{year}-{last:04d}"


def assigned_role_for_stage(doc_type: str, stage: str) -> Optional[str]:
    candidates = []
    for _aid, meta in ACTIONS.items():
        if doc_type in meta.get("types", []) and stage in meta.get("from", []):
            candidates.extend(meta.get("roles", []))

    if not candidates:
        return None

    prefer = [Role.APPROVER, Role.EXECUTOR, Role.PDZK, Role.CLERK, Role.AUDITOR]
    for r in prefer:
        if r in candidates:
            return r
    return candidates[0]


def set_assignment_after_transition(d: Document, meta: dict):
    if is_terminal(d.stage) or d.stage == "deleted":
        set_assigned_role(d, None)
        return

    next_role = meta.get("to_role")
    if next_role is None:
        next_role = assigned_role_for_stage(d.doc_type, d.stage)

    set_assigned_role(d, next_role)


# -------------------- Роуты --------------------

@main_bp.route("/users")
@login_required
@roles_required(Role.AUDITOR, Role.APPROVER)
def users():
    users_list = User.query.order_by(User.id.asc()).all()
    return render_template(
        "users.html",
        me=current_user,
        users=users_list,
        ROLE_TITLES=ROLE_TITLES,
    )


@main_bp.route("/journal")
@login_required
def journal():
    if current_user.role not in [Role.AUDITOR, Role.APPROVER]:
        abort(403)

    q = (request.args.get("q") or "").strip()

    events_q = (
        AuditEvent.query
        .join(Document, AuditEvent.document_id == Document.id)
    )

    if q:
        like = f"%{q}%"
        events_q = events_q.filter(
            or_(
                Document.title.ilike(like),
                Document.reg_number.ilike(like)
            )
        )

    events = (
        events_q
        .order_by(AuditEvent.document_id.desc(), AuditEvent.created_at.asc())
        .all()
    )

    grouped = {}
    for e in events:
        doc = getattr(e, "document", None)
        if not doc:
            continue

        actor = getattr(e, "actor", None)
        actor_name = getattr(actor, "username", None) or getattr(actor, "login", None) or f"ID {e.actor_id}"
        actor_role_raw = getattr(actor, "role", None) if actor else None
        actor_role = ROLE_TITLES.get(actor_role_raw, actor_role_raw or "")

        g = grouped.setdefault(doc.id, {"doc": doc, "events": []})
        g["events"].append({
            "action": e.action,
            "at": e.created_at,
            "from": STAGE_LABELS.get(e.from_stage, e.from_stage) if e.from_stage else "—",
            "to": STAGE_LABELS.get(e.to_stage, e.to_stage) if e.to_stage else "—",
            "actor_name": actor_name,
            "actor_role": actor_role,
        })

    docs_log = []
    for _doc_id, data in grouped.items():
        evs = data["events"]
        created = evs[0] if evs else None
        rest = evs[1:] if len(evs) > 1 else []
        last = evs[-1] if evs else None

        docs_log.append({
            "doc": data["doc"],
            "created": created,
            "rest": rest,
            "last": last,
            "count": len(evs),
        })

    docs_log.sort(key=lambda x: x["doc"].created_at or datetime.min, reverse=True)

    return render_template(
        "audit_journal.html",
        me=current_user,
        q=q,
        docs_log=docs_log,
        DOC_TYPE_LABELS=DOC_TYPE_LABELS,
        STAGE_LABELS=STAGE_LABELS,
    )


@main_bp.route("/")
@login_required
def index():
    active_docs = (
        Document.query
        .filter(Document.stage != "deleted")
        .order_by(Document.created_at.desc())
        .all()
    )

    for d in active_docs:
        attach_runtime_assigned(d)

    inbox_docs = []
    for d in active_docs:
        if d.stage in ("archived", "destroyed", "declassified"):
            continue
        assigned = getattr(d, "assigned_to_role", None)
        if assigned and assigned == current_user.role:
            inbox_docs.append(d)

    return render_template(
        "index.html",
        inbox_docs=inbox_docs,
        docs=active_docs,
        me=current_user,
        DOC_TYPE_LABELS=DOC_TYPE_LABELS,
        STAGE_LABELS=STAGE_LABELS
    )


@main_bp.route("/docs/new", methods=["GET", "POST"])
@login_required
@roles_required(Role.CLERK)
def new_doc():
    if request.method == "GET":
        return render_template("new_doc.html", me=current_user, DOC_TYPE_LABELS=DOC_TYPE_LABELS)

    doc_type = request.form.get("doc_type")
    title = request.form.get("title", "").strip()

    if doc_type not in DOC_TYPE_LABELS or not title:
        return render_template(
            "new_doc.html",
            me=current_user,
            DOC_TYPE_LABELS=DOC_TYPE_LABELS,
            error="Заполни тип и название"
        )

    payload = {
        "org_name": request.form.get("org_name", "").strip(),
        "sign_pos": request.form.get("sign_pos", "").strip(),
        "sign_name": request.form.get("sign_name", "").strip(),
    }

    conf = request.form.get("confidentiality", "").strip()

    if doc_type == DocumentType.LETTER:
        payload["addressee"] = request.form.get("letter_addressee", "").strip()
        payload["body"] = request.form.get("letter_body", "").strip()

    elif doc_type == DocumentType.INSTRUCTION:
        payload["purpose"] = request.form.get("ins_purpose", "").strip()
        payload["scope"] = request.form.get("ins_scope", "").strip()
        payload["steps_raw"] = request.form.get("ins_steps", "").strip()
        payload["responsibility"] = request.form.get("ins_responsibility", "").strip()

    elif doc_type == DocumentType.PACKET:
        payload["list_raw"] = request.form.get("pack_list", "").strip()
        payload["note"] = request.form.get("pack_note", "").strip()

    st = initial_stage(doc_type)
    payload["_assigned_to_role"] = assigned_role_for_stage(doc_type, st)

    d = Document(
        doc_type=doc_type,
        title=title,
        reg_number=make_reg_number(doc_type),
        stage=st,
        created_by_id=current_user.id,
        content_json=json.dumps(payload, ensure_ascii=False),
        confidentiality=conf or None,
    )
    db.session.add(d)
    db.session.flush()

    db.session.add(AuditEvent(
        document_id=d.id,
        actor_id=current_user.id,
        action="Создан документ",
        from_stage=None,
        to_stage=d.stage
    ))
    db.session.commit()

    return redirect(url_for("main.doc_detail", doc_id=d.id))


@main_bp.route("/docs/<int:doc_id>")
@login_required
def doc_detail(doc_id: int):
    d = Document.query.get_or_404(doc_id)

    if d.stage == "deleted":
        abort(404)

    if d.stage == "destroyed":
        return redirect(url_for("main.index"))

    events = (
        AuditEvent.query
        .filter_by(document_id=d.id)
        .order_by(AuditEvent.created_at.desc())
        .all()
    )

    progress, step_now, step_total, stage_label = calc_progress(d.doc_type, d.stage)
    payload = load_payload(d)

    if not payload.get("_assigned_to_role") and not is_terminal(d.stage) and d.stage != "deleted":
        payload["_assigned_to_role"] = assigned_role_for_stage(d.doc_type, d.stage)
        save_payload(d, payload)
        db.session.commit()

    attach_runtime_assigned(d, payload)
    actions = available_actions(d.doc_type, d.stage, current_user.role, getattr(d, "assigned_to_role", None))

    ins_steps = [x.strip() for x in (payload.get("steps_raw") or "").splitlines() if x.strip()]
    pack_items = [x.strip("-• \t") for x in (payload.get("list_raw") or "").splitlines() if x.strip()]

    history = []
    for e in events:
        actor = getattr(e, "actor", None)
        actor_name = getattr(actor, "username", None) or getattr(actor, "login", None) or f"ID {e.actor_id}"
        actor_role_raw = getattr(actor, "role", None) if actor else None
        actor_role = ROLE_TITLES.get(actor_role_raw, actor_role_raw or "")

        history.append({
            "action": e.action,
            "at": e.created_at,
            "from": STAGE_LABELS.get(e.from_stage, e.from_stage) if e.from_stage else "—",
            "to": STAGE_LABELS.get(e.to_stage, e.to_stage) if e.to_stage else "—",
            "actor_name": actor_name,
            "actor_role": actor_role,
        })

    users_for_send = User.query.filter(User.id != current_user.id).order_by(User.username.asc()).all()

    return render_template(
        "doc_detail.html",
        me=current_user,
        d=d,
        DOC_TYPE_LABELS=DOC_TYPE_LABELS,
        STAGE_LABELS=STAGE_LABELS,
        ROLE_TITLES=ROLE_TITLES,
        progress=progress,
        step_now=step_now,
        step_total=step_total,
        stage_label=stage_label,
        actions=actions,
        payload=payload,
        ins_steps=ins_steps,
        pack_items=pack_items,
        history=history,
        users_for_send=users_for_send,
    )


@main_bp.route("/docs/<int:doc_id>/edit", methods=["GET", "POST"])
@login_required
def edit_doc(doc_id: int):
    d = Document.query.get_or_404(doc_id)

    if d.stage in ("deleted", "destroyed", "archived", "declassified"):
        abort(403)

    if d.created_by_id != current_user.id:
        abort(403)

    payload = load_payload(d)

    if request.method == "GET":
        return render_template(
            "edit_doc.html",
            me=current_user,
            d=d,
            payload=payload,
            DOC_TYPE_LABELS=DOC_TYPE_LABELS
        )

    title = request.form.get("title", "").strip()
    if not title:
        return render_template(
            "edit_doc.html",
            me=current_user,
            d=d,
            payload=payload,
            DOC_TYPE_LABELS=DOC_TYPE_LABELS,
            error="Название не может быть пустым"
        )

    old_stage = d.stage

    new_payload = {
        "org_name": request.form.get("org_name", "").strip(),
        "sign_pos": request.form.get("sign_pos", "").strip(),
        "sign_name": request.form.get("sign_name", "").strip(),
    }

    conf = request.form.get("confidentiality", "").strip()
    d.confidentiality = conf or None

    if d.doc_type == DocumentType.LETTER:
        new_payload["addressee"] = request.form.get("letter_addressee", "").strip()
        new_payload["body"] = request.form.get("letter_body", "").strip()

    elif d.doc_type == DocumentType.INSTRUCTION:
        new_payload["purpose"] = request.form.get("ins_purpose", "").strip()
        new_payload["scope"] = request.form.get("ins_scope", "").strip()
        new_payload["steps_raw"] = request.form.get("ins_steps", "").strip()
        new_payload["responsibility"] = request.form.get("ins_responsibility", "").strip()

    elif d.doc_type == DocumentType.PACKET:
        new_payload["list_raw"] = request.form.get("pack_list", "").strip()
        new_payload["note"] = request.form.get("pack_note", "").strip()

    d.title = title

    init_st = initial_stage(d.doc_type)

    if old_stage != init_st:
        d.stage = init_st
        new_payload["_assigned_to_role"] = assigned_role_for_stage(d.doc_type, d.stage)
        d.content_json = json.dumps(new_payload, ensure_ascii=False)

        e = AuditEvent(
            document_id=d.id,
            actor_id=current_user.id,
            action="Документ отредактирован. Процесс сброшен и запущен заново",
            from_stage=old_stage,
            to_stage=d.stage
        )
    else:
        new_payload["_assigned_to_role"] = assigned_role_for_stage(d.doc_type, d.stage)
        d.content_json = json.dumps(new_payload, ensure_ascii=False)

        e = AuditEvent(
            document_id=d.id,
            actor_id=current_user.id,
            action="Документ отредактирован",
            from_stage=old_stage,
            to_stage=old_stage
        )

    db.session.add(e)
    db.session.commit()
    return redirect(url_for("main.doc_detail", doc_id=d.id))


@main_bp.post("/docs/<int:doc_id>/delete")
@login_required
@roles_required(Role.AUDITOR)
def doc_delete(doc_id: int):
    d = Document.query.get_or_404(doc_id)

    if d.stage == "destroyed":
        return redirect(url_for("main.index"))

    from_stage = d.stage
    d.stage = "deleted"

    payload = load_payload(d)
    payload.pop("_assigned_to_role", None)
    save_payload(d, payload)

    db.session.add(AuditEvent(
        document_id=d.id,
        actor_id=current_user.id,
        action="Удалён наблюдателем",
        from_stage=from_stage,
        to_stage="deleted",
    ))
    db.session.commit()

    return redirect(url_for("main.index"))


@main_bp.route("/docs/<int:doc_id>/action/<action_id>", methods=["POST"])
@login_required
def doc_action(doc_id: int, action_id: str):
    d = Document.query.get_or_404(doc_id)
    meta = ACTIONS.get(action_id)
    if not meta:
        abort(404)

    if d.doc_type not in meta.get("types", []):
        abort(403)

    if d.stage not in meta.get("from", []):
        abort(403)

    if current_user.role not in meta.get("roles", []):
        abort(403)

    payload = load_payload(d)
    assigned_to = payload.get("_assigned_to_role")
    if assigned_to and assigned_to != current_user.role:
        abort(403)

    from_stage = d.stage
    d.stage = meta.get("to")

    set_assignment_after_transition(d, meta)

    db.session.add(AuditEvent(
        document_id=d.id,
        actor_id=current_user.id,
        action=meta.get("label", action_id),
        from_stage=from_stage,
        to_stage=d.stage
    ))
    db.session.commit()

    if d.stage == "destroyed":
        return redirect(url_for("main.index"))

    return redirect(url_for("main.doc_detail", doc_id=d.id))


@main_bp.route("/docs/<int:doc_id>/download/docx")
@login_required
def download_docx(doc_id: int):
    d = Document.query.get_or_404(doc_id)
    payload = load_payload(d)

    doc = DocxDocument()

    if d.confidentiality:
        doc.add_paragraph(d.confidentiality)

    org = payload.get("org_name") or ""
    if org:
        doc.add_paragraph(org)

    doc.add_paragraph(f"{DOC_TYPE_LABELS.get(d.doc_type, d.doc_type).upper()} № {d.reg_number}")
    doc.add_paragraph(d.title)
    doc.add_paragraph("")

    if d.doc_type == DocumentType.LETTER:
        addressee = payload.get("addressee", "")
        body = payload.get("body", "")

        if addressee:
            doc.add_paragraph("Кому:")
            doc.add_paragraph(addressee)
            doc.add_paragraph("")
        if body:
            doc.add_paragraph("Текст письма:")
            doc.add_paragraph(body)

    elif d.doc_type == DocumentType.INSTRUCTION:
        purpose = payload.get("purpose", "")
        scope = payload.get("scope", "")
        steps_raw = payload.get("steps_raw", "")
        resp = payload.get("responsibility", "")

        if purpose:
            doc.add_paragraph("1. Цель")
            doc.add_paragraph(purpose)
        if scope:
            doc.add_paragraph("2. Область применения")
            doc.add_paragraph(scope)

        steps = [x.strip() for x in steps_raw.splitlines() if x.strip()]
        if steps:
            doc.add_paragraph("3. Порядок действий")
            for i, line in enumerate(steps, 1):
                doc.add_paragraph(f"{i}. {line}")

        if resp:
            doc.add_paragraph("4. Ответственность")
            doc.add_paragraph(resp)

    elif d.doc_type == DocumentType.PACKET:
        lst_raw = payload.get("list_raw", "")
        note = payload.get("note", "")

        items = [x.strip("-• \t") for x in lst_raw.splitlines() if x.strip()]
        doc.add_paragraph("Состав пакета:")
        for x in items:
            doc.add_paragraph(f"- {x}")

        if note:
            doc.add_paragraph("")
            doc.add_paragraph("Примечание:")
            doc.add_paragraph(note)

    sign_pos = payload.get("sign_pos") or ""
    sign_name = payload.get("sign_name") or ""
    if sign_pos or sign_name:
        doc.add_paragraph("")
        doc.add_paragraph(f"{sign_pos} ____________________ {sign_name}".strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{d.reg_number}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@main_bp.post("/docs/<int:doc_id>/secure/send")
@login_required
def secure_send(doc_id: int):
    d = Document.query.get_or_404(doc_id)

    if d.stage == "deleted":
        abort(404)

    to_user_id = int(request.form.get("to_user_id") or 0)
    to_user = User.query.get_or_404(to_user_id)

    sender_keys = ensure_user_keys(current_user)
    recipient_keys = ensure_user_keys(to_user)

    payload = load_payload(d)
    snapshot = build_snapshot(d, payload)

    h = streebog256(snapshot)

    d_priv = int(sender_keys.sign_d)
    sig = gost_sign(snapshot, d_priv)

    aad_obj = {"from": current_user.username, "to": to_user.username, "doc_id": d.id}
    aad = json.dumps(aad_obj, ensure_ascii=False, sort_keys=True).encode("utf-8")

    key_bytes = bytes.fromhex(recipient_keys.enc_key_hex)
    enc = encrypt_for_recipient(snapshot, key_bytes, aad)

    mac = kuz_cmac(key_bytes, b64d(enc["ct"]))

    envelope = {
        "v": 1,
        "meta": {
            "from_user": current_user.username,
            "to_user": to_user.username,
            "doc_ref": {"id": d.id, "reg_number": d.reg_number, "title": d.title},
        },
        "crypto": {
            "hash": {"alg": "streebog-256", "value": b64e(h)},
            "sign": {
                "alg": "gost-like-demo",
                "qx": sender_keys.sign_qx,
                "qy": sender_keys.sign_qy,
                "sig": b64e(sig),
            },
            "enc": enc,
            "imito": {"alg": "kuz-cmac-or-fallback", "value": b64e(mac)},
        },
    }

    m = SecureMessage(
        from_user_id=current_user.id,
        to_user_id=to_user.id,
        document_id=d.id,
        envelope_json=json.dumps(envelope, ensure_ascii=False),
        status="new",
    )
    db.session.add(m)

    db.session.add(AuditEvent(
        document_id=d.id,
        actor_id=current_user.id,
        action=f"Отправлен защищённый пакет пользователю {to_user.username}",
        from_stage=d.stage,
        to_stage=d.stage,
    ))
    db.session.commit()

    return redirect(url_for("main.doc_detail", doc_id=d.id))


@main_bp.route("/secure/inbox")
@login_required
def secure_inbox():
    msgs = (
        SecureMessage.query
        .filter_by(to_user_id=current_user.id)
        .order_by(SecureMessage.created_at.desc())
        .all()
    )
    return render_template("secure_inbox.html", me=current_user, msgs=msgs)


@main_bp.route("/secure/<int:msg_id>")
@login_required
def secure_view(msg_id: int):
    msg = SecureMessage.query.get_or_404(msg_id)
    if msg.to_user_id != current_user.id and msg.from_user_id != current_user.id:
        abort(403)

    env = json.loads(msg.envelope_json)

    # ключи получателя
    recipient_keys = ensure_user_keys(current_user)
    key_bytes = bytes.fromhex(recipient_keys.enc_key_hex)

    # расшифровка + проверка MGM-тега
    pt, ok_mgm = decrypt_for_recipient(env["crypto"]["enc"], key_bytes)

    # проверка имитовставки
    mac_expected_b64 = env.get("crypto", {}).get("imito", {}).get("value", "")
    ct_b64 = env.get("crypto", {}).get("enc", {}).get("ct", "")
    mac_real = kuz_cmac(key_bytes, b64decode(ct_b64)) if ct_b64 else b""
    ok_mac = (mac_real == base64.b64decode(mac_expected_b64)) if mac_expected_b64 else False

    # проверка хеша
    h_expected_b64 = env.get("crypto", {}).get("hash", {}).get("value", "")
    h_real = streebog256(pt) if pt else b""
    ok_hash = (h_real == base64.b64decode(h_expected_b64)) if h_expected_b64 else False

    # проверка подписи
    sig_b64 = env.get("crypto", {}).get("sign", {}).get("sig", "")
    qx = int(env.get("crypto", {}).get("sign", {}).get("qx", 0) or 0)
    qy = int(env.get("crypto", {}).get("sign", {}).get("qy", 0) or 0)
    sig = base64.b64decode(sig_b64) if sig_b64 else b""
    ok_sign = gost_verify(pt, sig, qx, qy) if pt and sig and qx and qy else False

    snapshot = None
    if ok_mgm and pt:
        try:
            snapshot = json.loads(pt.decode("utf-8"))
        except Exception:
            snapshot = None

    # пометим как opened
    if msg.status == "new" and msg.to_user_id == current_user.id:
        msg.status = "opened"
        msg.opened_at = datetime.utcnow()
        db.session.commit()

    def short_val(s: str, head: int = 14, tail: int = 14) -> str:
        if not s:
            return "—"
        if len(s) <= head + tail + 1:
            return s
        return f"{s[:head]}…{s[-tail:]}"

    # красивый “паспорт крипты”
    enc_obj = env.get("crypto", {}).get("enc", {}) or {}
    nonce_b64 = enc_obj.get("nonce") or enc_obj.get("iv") or ""
    tag_b64 = enc_obj.get("tag") or enc_obj.get("mac") or ""
    ct_len = 0
    try:
        ct_len = len(base64.b64decode(ct_b64)) if ct_b64 else 0
    except Exception:
        ct_len = 0

    crypto_view = {
        "encryption": {
            "title": "Шифрование",
            "alg": "Кузнечик (режим MGM)",
            "ok": ok_mgm,
            "details": [
                ("Шифртекст", f"{ct_len} байт"),
                ("Нонс", short_val(nonce_b64) if nonce_b64 else "—"),
                ("Тег MGM", short_val(tag_b64) if tag_b64 else "—"),
            ],
        },
        "mac": {
            "title": "Имитовставка",
            "alg": "CMAC-Кузнечик",
            "ok": ok_mac,
            "details": [
                ("Значение", short_val(mac_expected_b64)),
            ],
            "full": mac_expected_b64,
        },
        "hash": {
            "title": "Хэш",
            "alg": "Стрибог-256",
            "ok": ok_hash,
            "details": [
                ("Значение", short_val(h_expected_b64)),
            ],
            "full": h_expected_b64,
        },
        "sign": {
            "title": "ЭЦП",
            "alg": "ГОСТ 34.10 (подпись)",
            "ok": ok_sign,
            "details": [
                ("Подпись", short_val(sig_b64)),
                ("Открытый ключ", f"Q=({qx}, {qy})" if qx and qy else "—"),
            ],
            "full": sig_b64,
        },
    }

    # предпросмотр содержимого “человеческим языком”
    preview = None
    if snapshot:
        meta = snapshot.get("meta", {}) or {}
        payload = snapshot.get("payload", {}) or {}
        preview = {
            "doc_type": meta.get("doc_type"),
            "title": meta.get("title"),
            "reg_number": meta.get("reg_number"),
            "stage": meta.get("stage"),
            "created_at": meta.get("created_at"),
            "confidentiality": meta.get("confidentiality"),
            "payload": payload,
        }

    envelope_pretty = json.dumps(env, ensure_ascii=False, indent=2)
    snapshot_pretty = json.dumps(snapshot, ensure_ascii=False, indent=2) if snapshot else None

    return render_template(
        "secure_view.html",
        me=current_user,
        msg=msg,
        env=env,
        preview=preview,
        crypto_view=crypto_view,
        ok_mgm=ok_mgm,
        ok_mac=ok_mac,
        ok_hash=ok_hash,
        ok_sign=ok_sign,
        envelope_pretty=envelope_pretty,
        snapshot_pretty=snapshot_pretty,
    )



@main_bp.post("/secure/<int:msg_id>/accept")
@login_required
def secure_accept(msg_id: int):
    msg = SecureMessage.query.get_or_404(msg_id)
    if msg.to_user_id != current_user.id:
        abort(403)

    env = json.loads(msg.envelope_json)

    recipient_keys = ensure_user_keys(current_user)
    key_bytes = bytes.fromhex(recipient_keys.enc_key_hex)

    pt, ok_mgm = decrypt_for_recipient(env["crypto"]["enc"], key_bytes)
    if not ok_mgm:
        abort(400)

    snap = json.loads(pt.decode("utf-8"))
    meta = snap.get("meta", {})
    payload = snap.get("payload", {})

    doc_type = meta.get("doc_type")
    title = meta.get("title") or "(без названия)"

    st = initial_stage(doc_type)

    # важно: назначение в JSON, чтобы не требовать колонки assigned_to_role
    payload["_assigned_to_role"] = assigned_role_for_stage(doc_type, st)

    d = Document(
        doc_type=doc_type,
        title=title,
        reg_number=make_reg_number(doc_type),
        stage=st,
        created_by_id=current_user.id,
        content_json=json.dumps(payload, ensure_ascii=False),
        confidentiality=meta.get("confidentiality") or None,
    )
    db.session.add(d)
    db.session.flush()

    db.session.add(AuditEvent(
        document_id=d.id,
        actor_id=current_user.id,
        action=f"Принят защищённый пакет (сообщение #{msg.id})",
        from_stage=None,
        to_stage=d.stage
    ))

    msg.status = "accepted"
    db.session.commit()

    return redirect(url_for("main.doc_detail", doc_id=d.id))
